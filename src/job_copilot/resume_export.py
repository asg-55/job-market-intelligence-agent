from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

INK = RGBColor(16, 35, 31)
ACCENT = RGBColor(46, 116, 181)
MUTED = RGBColor(92, 108, 102)


def build_resume_docx(resume: dict) -> bytes:
    """Create an editable Word export without changing the stored resume."""
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    _configure_styles(document)
    document.core_properties.author = ""
    document.core_properties.last_modified_by = ""
    document.core_properties.title = resume["name"]

    document.add_paragraph(resume["name"], style="Resume Title")

    roles = ", ".join(resume.get("target_roles", []))
    if roles:
        document.add_paragraph(roles, style="Resume Subtitle")

    for raw_line in resume["content"].splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if _looks_like_heading(line):
            document.add_paragraph(line.rstrip(":"), style="Heading 1")
        elif line.startswith(("- ", "* ", "• ")):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        else:
            document.add_paragraph(line)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _configure_styles(document) -> None:
    normal = document.styles["Normal"]
    _style_font(normal, size=11, color=INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    title = document.styles.add_style("Resume Title", WD_STYLE_TYPE.PARAGRAPH)
    _style_font(title, size=24, color=INK, bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(4)
    title.paragraph_format.keep_with_next = True

    subtitle = document.styles.add_style("Resume Subtitle", WD_STYLE_TYPE.PARAGRAPH)
    _style_font(subtitle, size=11, color=MUTED)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(16)
    subtitle.paragraph_format.keep_with_next = True

    heading = document.styles["Heading 1"]
    _style_font(heading, size=16, color=ACCENT, bold=True)
    heading.paragraph_format.space_before = Pt(16)
    heading.paragraph_format.space_after = Pt(8)
    heading.paragraph_format.keep_with_next = True

    bullet = document.styles["List Bullet"]
    _style_font(bullet, size=11, color=INK)
    bullet.paragraph_format.left_indent = Inches(0.5)
    bullet.paragraph_format.first_line_indent = Inches(-0.25)
    bullet.paragraph_format.space_after = Pt(8)
    bullet.paragraph_format.line_spacing = 1.167
    bullet.paragraph_format.keep_together = True


def _style_font(style, *, size: int, color: RGBColor, bold: bool = False) -> None:
    style.font.name = "Calibri"
    style._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), "Calibri")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    style.font.size = Pt(size)
    style.font.color.rgb = color
    style.font.bold = bold


def _looks_like_heading(line: str) -> bool:
    if len(line) > 70:
        return False
    words = line.rstrip(":").split()
    return line.endswith(":") or (len(words) <= 6 and line.isupper())
